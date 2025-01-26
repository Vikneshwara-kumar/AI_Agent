import streamlit as st
import os
from groq import Groq
from typing import List, Dict, Optional
import time
from dataclasses import dataclass
from enum import Enum
from docx import Document
from docx.shared import Pt
from io import BytesIO

def convert_md_to_docx(md_text: str) -> BytesIO:
    """Convert markdown text to structured Word document"""
    doc = Document()
    
    # Add title style
    style = doc.styles['Title']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(24)
    
    # Add heading styles
    for level in range(1, 4):
        style = doc.styles[f'Heading {level}']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(24 - (level * 2))
        font.bold = True
    
    # Parse markdown content
    current_heading = None
    for line in md_text.split('\n'):
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line:
            p = doc.add_paragraph(line)
            p.style = 'BodyText'
    
    # Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

class PRDSection(Enum):
    INTRODUCTION = "Introduction"
    USER_STORIES = "User Stories"
    TECHNICAL_REQUIREMENTS = "Technical Requirements"
    BUDGET = "Budget"
    TIMELINE = "Timeline"
    MARKET_ANALYSIS = "Market Analysis"

@dataclass
class PRDContent:
    sections: Dict[PRDSection, str]
    original_text: str

def initialize_groq_client() -> Optional[Groq]:
    """Initialize Groq client using API key from environment or Streamlit secrets."""
    try:
        if 'GROQ_API_KEY' in st.secrets:
            return Groq(api_key=st.secrets['GROQ_API_KEY'])
        elif 'GROQ_API_KEY' in os.environ:
            return Groq(api_key=os.environ['GROQ_API_KEY'])
        return None
    except Exception as e:
        st.error(f"Error initializing Groq client: {str(e)}")
        return None

def parse_prd_sections(text: str, client: Optional[Groq]) -> PRDContent:
    """Parse PRD text into relevant sections using AI."""
    sections = {}
    
    if client is None:
        st.error("No Groq client available for parsing")
        return PRDContent(sections={}, original_text=text)
    
    prompt = f"""Parse this PRD into the following sections:
    - Introduction
    - User Stories
    - Technical Requirements
    - Budget
    - Timeline
    - Market Analysis

    PRD Text: {text}
    
    Return the content organized by these sections."""
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="gemma2-9b-it",
            temperature=0.3,
            max_tokens=2000
        )
        parsed_content = response.choices[0].message.content
        
        # Simple parsing logic - can be enhanced with better section detection
        current_section = None
        current_content = []
        
        for line in parsed_content.split('\n'):
            for section in PRDSection:
                if section.value in line:
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = section
                    current_content = []
                    break
            else:
                if current_section:
                    current_content.append(line)
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
            
        return PRDContent(sections=sections, original_text=text)
    except Exception as e:
        st.error(f"Error parsing PRD: {str(e)}")
        return PRDContent(sections={}, original_text=text)

class Agent:
    def __init__(self, role: str, expertise: List[str], focus_sections: List[PRDSection], client: Optional[Groq]):
        self.role = role
        self.expertise = expertise
        self.focus_sections = focus_sections
        self.client = client
    
    def analyze_section(self, section: PRDSection, content: str) -> str:
        """Analyze a specific PRD section based on agent's expertise."""
        if self.client is None:
            return "Error: No Groq client available for analysis"
        
        prompt = f"""As a {self.role} with expertise in {', '.join(self.expertise)}, 
        analyze this {section.value} section:
        
        Content: {content}
        
        Provide specific feedback focusing on:
        1. Completeness and clarity
        2. Areas of improvement related to your expertise
        3. Potential risks or challenges
        4. Specific recommendations
        5. Impact on other areas of the project"""
        
        try:
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="gemma2-9b-it",
                temperature=0.3,
                max_tokens=1000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating feedback: {str(e)}"
    
    def debate_feedback(self, other_feedback: Dict[str, str]) -> str:
        """Generate a response to other agents' feedback."""
        if self.client is None:
            return "Error: No Groq client available for debate"
            
        truncated_feedback = "\n".join([f"{role}: {feedback[:200]}..." 
                                      for role, feedback in other_feedback.items()])
        
        prompt = f"""As a {self.role}, review and respond to this feedback:
        
        {truncated_feedback}
        
        Provide:
        1. Points of agreement
        2. Constructive disagreements
        3. Additional considerations from your expertise
        4. Suggested compromises for conflicting viewpoints"""
        
        try:
            response = self.client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="gemma2-9b-it",
                temperature=0.3,
                max_tokens=800
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating debate response: {str(e)}"

class Orchestrator:
    def __init__(self, client: Optional[Groq]):
        self.client = client
        self.agents = {
            "UX Lead": Agent("UX Lead", 
                           ["user experience", "interaction design", "accessibility"],
                           [PRDSection.USER_STORIES, PRDSection.TECHNICAL_REQUIREMENTS],
                           client),
            
            "Data Scientist": Agent("Data Scientist",
                                  ["data requirements", "analytics", "technical feasibility"],
                                  [PRDSection.TECHNICAL_REQUIREMENTS],
                                  client),
            
            "Software Engineer": Agent("Software Engineer",
                                     ["technical implementation", "architecture", "security"],
                                     [PRDSection.TECHNICAL_REQUIREMENTS, PRDSection.TIMELINE],
                                     client),
            
            "DevOps Engineer": Agent("DevOps Engineer",
                                   ["deployment", "scalability", "infrastructure"],
                                   [PRDSection.TECHNICAL_REQUIREMENTS, PRDSection.TIMELINE],
                                   client),
            
            "Marketing Director": Agent("Marketing Director",
                                      ["market positioning", "user needs", "branding"],
                                      [PRDSection.MARKET_ANALYSIS, PRDSection.USER_STORIES],
                                      client),
            
            "Finance Manager": Agent("Finance Manager",
                                   ["cost analysis", "budgeting", "ROI"],
                                   [PRDSection.BUDGET, PRDSection.TIMELINE],
                                   client),
            
            "Project Manager": Agent("Project Manager",
                                   ["project planning", "resource allocation", "risk management"],
                                   [PRDSection.TIMELINE, PRDSection.BUDGET],
                                   client),
            
            "Product Strategy Expert": Agent("Product Strategy Expert",
                                           ["product vision", "market trends", "competitive analysis"],
                                           [PRDSection.INTRODUCTION, PRDSection.MARKET_ANALYSIS],
                                           client)
        }
    
    def analyze_prd(self, prd_content: PRDContent) -> Dict[str, Dict[PRDSection, str]]:
        """Generate section-specific feedback from each agent."""
        feedback = {agent_role: {} for agent_role in self.agents.keys()}
        
        with st.status("Analyzing PRD sections...", expanded=True) as status:
            for section in PRDSection:
                if section in prd_content.sections:
                    st.write(f"Analyzing {section.value}...")
                    for role, agent in self.agents.items():
                        if section in agent.focus_sections:
                            feedback[role][section] = agent.analyze_section(
                                section, 
                                prd_content.sections[section]
                            )
                            time.sleep(1)  # Rate limiting
        
        return feedback
    
    def facilitate_debate(self, feedback: Dict[str, Dict[PRDSection, str]]) -> str:
        """Facilitate debate between agents on their feedback."""
        if self.client is None:
            return "Error: No Groq client available for debate"
            
        debate_rounds = []
        
        with st.status("Facilitating debate...", expanded=True) as status:
            for section in PRDSection:
                section_feedback = {
                    role: agent_feedback.get(section, "")
                    for role, agent_feedback in feedback.items()
                }
                
                if any(section_feedback.values()):
                    st.write(f"Debating {section.value}...")
                    for role, agent in self.agents.items():
                        if section in agent.focus_sections:
                            response = agent.debate_feedback(section_feedback)
                            debate_rounds.append(f"{role} on {section.value}:\n{response}")
                            time.sleep(1)  # Rate limiting
        
        return "\n\n".join(debate_rounds)
    
    def generate_improved_prd(self, original_prd: PRDContent, 
                            feedback: Dict[str, Dict[PRDSection, str]], 
                            debate: str) -> str:
        """Generate improved PRD based on feedback and debate."""
        if self.client is None:
            return "Error: No Groq client available for generating improved PRD"
            
        sections_summary = {}
        
        for section in PRDSection:
            section_feedback = "\n".join([
                f"{role}: {feedback[role].get(section, '')}"
                for role in self.agents.keys()
                if section in self.agents[role].focus_sections
            ])
            
            prompt = f"""Improve this PRD section based on feedback:
            
            Original {section.value}:
            {original_prd.sections.get(section, '')}
            
            Feedback:
            {section_feedback}
            
            Generate an improved version incorporating the feedback."""
            
            try:
                response = self.client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="gemma2-9b-it",
                    temperature=0.3,
                    max_tokens=1000
                )
                sections_summary[section] = response.choices[0].message.content
            except Exception as e:
                sections_summary[section] = f"Error improving section: {str(e)}"
            
            time.sleep(1)  # Rate limiting
        
        return "\n\n".join([f"# {section.value}\n{content}" 
                           for section, content in sections_summary.items()])

def main():
    st.title("PRD Improviser🚀")
    st.write("This app helps you analyze and enhance your PRD by leveraging AI-powered agents to provide Stakeholder's feedback, analysis, and facilitate discussions. Upload your PRD for multi-agent analysis and improvement suggestions.")

    # Initialize client
    client = initialize_groq_client()
    if client is None:
        st.sidebar.warning("No Groq API key found.")
        api_key = st.sidebar.text_input("Enter your Groq API Key:", type="password")
        if api_key:
            try:
                client = Groq(api_key=api_key)
                st.sidebar.success("Successfully initialized!")
            except Exception as e:
                st.sidebar.error(f"Error: {str(e)}")
    
    if client is None:
        st.warning("Please provide a valid Groq API key to continue.")
        st.stop()

    # Main interface
    prd_text = st.text_area("Paste your PRD text here:", height=200)
    
    if st.button("Analyze PRD") and prd_text:
        orchestrator = Orchestrator(client)
        
        # Create tabs for workflow steps - CHANGED ORDER HERE
        final_tab, parse_tab, analysis_tab, debate_tab = st.tabs([
            "4. Improved PRD",  # Now first
            "1. Parse PRD",
            "2. Analysis",
            "3. Debate"
        ])
        
        # Step 1: Parse PRD
        with parse_tab:
            st.header("PRD Sections")
            prd_content = parse_prd_sections(prd_text, client)
            
            for section, content in prd_content.sections.items():
                with st.expander(section.value):
                    st.write(content) 
        
        # Step 2: Generate Analysis
        with analysis_tab:
            st.header("Agent Analysis")
            feedback = orchestrator.analyze_prd(prd_content)
            
            for section in PRDSection:
                if section in prd_content.sections:
                    with st.expander(f"{section.value} Analysis"):
                        for role, agent_feedback in feedback.items():
                            if section in agent_feedback:
                                st.subheader(role)
                                st.write(agent_feedback[section])
        
        # Step 3: Facilitate Debate
        with debate_tab:
            st.header("Agent Debate")
            debate = orchestrator.facilitate_debate(feedback)
            st.write(debate)
        
        # Step 4: Generate Improved PRD (Now first tab)
        with final_tab:
            st.header("Improved PRD")
            improved_prd = orchestrator.generate_improved_prd(
                prd_content,
                feedback,
                debate
            )

            # Convert to Word document
            docx_file = convert_md_to_docx(improved_prd)
            
            # Download button at the top
            st.download_button(
                "Download Improved PRD",
                data=docx_file,
                file_name="improved_prd.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            
            # Dropdown comparisons
            with st.expander("🔽 Compare Original vs Improved PRD", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.subheader("Original PRD")
                    st.write(prd_text)
                
                with col2:
                    st.subheader("Improved PRD")
                    st.write(improved_prd)
            
            # Optional feedback
            st.divider()
            st.subheader("Feedback")
            user_rating = st.slider("How helpful was this analysis?", 1, 5, 3)
            user_feedback = st.text_area("Additional feedback (optional):")
            if st.button("Submit Feedback"):
                st.success("Thank you for your feedback!")

if __name__ == "__main__":
    main()